from __future__ import annotations

import re
import zipfile
import xml.etree.ElementTree as ET
from typing import Dict, List

from docx import Document


def _strip_ns(tag: str) -> str:
    if "}" in tag:
        return tag.split("}", 1)[1]
    return tag


def _extract_docx_textbox_texts(file_path: str) -> List[str]:
    texts: List[str] = []
    with zipfile.ZipFile(file_path, "r") as zf:
        xml_members = [name for name in zf.namelist() if name.startswith("word/") and name.endswith(".xml")]
        for member in sorted(xml_members):
            try:
                root = ET.fromstring(zf.read(member))
            except Exception:
                continue

            # DOCX 텍스트 상자는 주로 w:txbxContent 아래에 존재
            for node in root.iter():
                tag = _strip_ns(node.tag).lower()

                # 1) WordprocessingML 텍스트상자
                if tag == "txbxcontent":
                    raw = "".join(node.itertext())
                    text = re.sub(r"\s+", " ", raw).strip()
                    if text:
                        texts.append(text)

                # 2) DrawingML 도형 텍스트 (a:txBody)
                elif tag == "txbody":
                    chunks: List[str] = []
                    for child in node.iter():
                        if _strip_ns(child.tag).lower() == "t":
                            val = (child.text or "").strip()
                            if val:
                                chunks.append(val)
                    text = re.sub(r"\s+", " ", " ".join(chunks)).strip()
                    if text:
                        texts.append(text)
    return texts


def _extract_docx_table_texts(doc: Document) -> List[str]:
    texts: List[str] = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = re.sub(r"\s+", " ", cell.text or "").strip()
                if cell_text:
                    texts.append(cell_text)
    return texts


def extract_docx_pages(file_path: str, chars_per_page: int = 2000) -> List[Dict[str, object]]:
    doc = Document(file_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    textbox_texts = _extract_docx_textbox_texts(file_path)
    table_texts = _extract_docx_table_texts(doc)

    # 문단/표/텍스트상자 텍스트를 합치되 중복은 제거
    merged_texts: List[str] = []
    seen = set()
    for item in paragraphs + table_texts + textbox_texts:
        normalized = re.sub(r"\s+", " ", item).strip()
        if not normalized or normalized in seen:
            continue
        seen.add(normalized)
        merged_texts.append(normalized)

    pages: List[Dict[str, object]] = []
    buffer: List[str] = []
    current_len = 0
    page_no = 1

    for text in merged_texts:
        if current_len + len(text) > chars_per_page and buffer:
            pages.append({"page": page_no, "text": "\n".join(buffer).strip()})
            page_no += 1
            buffer = []
            current_len = 0

        buffer.append(text)
        current_len += len(text)

    if buffer:
        pages.append({"page": page_no, "text": "\n".join(buffer).strip()})

    return pages
